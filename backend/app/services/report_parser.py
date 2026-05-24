from __future__ import annotations

import subprocess
import tempfile
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile, ZipFile

from docx import Document
from PIL import Image


SUPPORTED_REPORT_EXTENSIONS = {".docx", ".doc", ".pdf"}


@dataclass(frozen=True)
class ParsedReport:
    text: str
    image_paths: list[Path]


def is_supported_report_file(filename: str) -> bool:
    return Path(filename).suffix.lower() in SUPPORTED_REPORT_EXTENSIONS


def validate_report_file(file_path: Path | str) -> tuple[bool, str]:
    path = Path(file_path)
    if path.suffix.lower() not in SUPPORTED_REPORT_EXTENSIONS:
        return False, "当前版本支持 .docx、.doc 和 .pdf 文件。"
    if not path.exists():
        return False, "文件不存在。"
    if path.stat().st_size <= 0:
        return False, "文件为空。"
    return True, ""


def validate_docx_file(file_path: Path | str) -> tuple[bool, str]:
    path = Path(file_path)
    if path.suffix.lower() != ".docx":
        return False, "当前文件不是 .docx 文件。"
    if not path.exists():
        return False, "文件不存在。"
    if path.stat().st_size <= 0:
        return False, "文件为空。"
    return True, ""


def parse_report_text(file_path: str | Path) -> str:
    """Extract text from docx/doc/pdf reports."""
    path = Path(file_path)
    ok, message = validate_report_file(path)
    if not ok:
        raise ValueError(message)

    suffix = path.suffix.lower()
    if suffix == ".docx":
        return parse_docx_text(path)
    if suffix == ".doc":
        with _converted_doc_to_docx(path) as docx_path:
            return parse_docx_text(docx_path)
    if suffix == ".pdf":
        return parse_pdf_text(path)
    raise ValueError("不支持的报告文件格式。")


def extract_report_images(file_path: str | Path, output_dir: str | Path) -> list[str]:
    """Extract or render images from docx/doc/pdf reports.

    - docx/doc: extract embedded images from Word media.
    - pdf: render each page to an image so teachers and AI can inspect pages.
    """
    path = Path(file_path)
    ok, message = validate_report_file(path)
    if not ok:
        raise ValueError(message)

    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx_images(path, output_dir)
    if suffix == ".doc":
        with _converted_doc_to_docx(path) as docx_path:
            return extract_docx_images(docx_path, output_dir)
    if suffix == ".pdf":
        return extract_pdf_page_images(path, output_dir)
    raise ValueError("不支持的报告文件格式。")


def parse_docx_text(file_path: str | Path) -> str:
    path = Path(file_path)
    ok, message = validate_docx_file(path)
    if not ok:
        raise ValueError(message)

    try:
        document = Document(str(path))
    except Exception as exc:
        raise ValueError(f"无法读取 Word 文件：{exc}") from exc

    sections: list[str] = []
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    if paragraphs:
        sections.append("【段落文本】")
        sections.extend(paragraphs)

    for table_index, table in enumerate(document.tables, start=1):
        table_lines: list[str] = []
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            if any(cells):
                table_lines.append(" | ".join(cells))
        if table_lines:
            sections.append(f"【表格{table_index}】")
            sections.extend(table_lines)

    return "\n".join(sections).strip()


def extract_docx_images(file_path: str | Path, output_dir: str | Path) -> list[str]:
    path = Path(file_path)
    ok, message = validate_docx_file(path)
    if not ok:
        raise ValueError(message)

    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    image_paths: list[str] = []
    try:
        with ZipFile(path) as archive:
            media_names = [
                name
                for name in archive.namelist()
                if name.startswith("word/media/") and not name.endswith("/")
            ]
            for index, media_name in enumerate(media_names, start=1):
                raw_bytes = archive.read(media_name)
                target = output_path / f"screenshot_{index:03d}.png"
                try:
                    image = Image.open(BytesIO(raw_bytes))
                    image.save(target, format="PNG")
                except Exception:
                    target.write_bytes(raw_bytes)
                image_paths.append(str(target))
    except BadZipFile as exc:
        raise ValueError("docx 文件结构损坏，无法提取图片。") from exc

    return image_paths


def parse_pdf_text(file_path: str | Path) -> str:
    path = Path(file_path)
    fitz = _try_load_pymupdf()
    if fitz is None:
        return _parse_pdf_text_with_pypdf(path)

    try:
        document = fitz.open(str(path))
    except Exception as exc:
        raise ValueError(f"无法读取 PDF 文件：{exc}") from exc

    sections: list[str] = []
    try:
        for page_index, page in enumerate(document, start=1):
            text = (page.get_text("text") or "").strip()
            if text:
                sections.append(f"【PDF第{page_index}页】")
                sections.append(text)
    finally:
        document.close()
    return "\n\n".join(sections).strip()


def extract_pdf_page_images(file_path: str | Path, output_dir: str | Path) -> list[str]:
    fitz = _try_load_pymupdf()
    if fitz is None:
        # Text-only PDF parsing is still useful. Page image rendering requires
        # PyMuPDF; when it is unavailable, keep parsing successful and let the
        # review page preview the original PDF directly.
        return []
    path = Path(file_path)
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    image_paths: list[str] = []
    try:
        document = fitz.open(str(path))
    except Exception as exc:
        raise ValueError(f"无法读取 PDF 文件：{exc}") from exc

    try:
        zoom = 1.6
        matrix = fitz.Matrix(zoom, zoom)
        for page_index, page in enumerate(document, start=1):
            target = output_path / f"screenshot_{page_index:03d}.png"
            pixmap = page.get_pixmap(matrix=matrix, alpha=False)
            pixmap.save(str(target))
            image_paths.append(str(target))
    finally:
        document.close()
    return image_paths


def extract_docx_report(docx_path: Path | str, image_output_dir: Path | str) -> ParsedReport:
    text = parse_docx_text(docx_path)
    image_paths = [Path(path) for path in extract_docx_images(docx_path, image_output_dir)]
    return ParsedReport(text=text, image_paths=image_paths)


def extract_report(report_path: Path | str, image_output_dir: Path | str) -> ParsedReport:
    text = parse_report_text(report_path)
    image_paths = [Path(path) for path in extract_report_images(report_path, image_output_dir)]
    return ParsedReport(text=text, image_paths=image_paths)


class _ConvertedDocContext:
    def __init__(self, source_path: Path):
        self.source_path = source_path
        self._temp_dir: tempfile.TemporaryDirectory[str] | None = None
        self.docx_path: Path | None = None

    def __enter__(self) -> Path:
        converter = _find_office_converter()
        if not converter:
            raise ValueError("解析 .doc 文件需要 LibreOffice 转换器。未检测到可用的 LibreOffice。")

        self._temp_dir = tempfile.TemporaryDirectory(prefix="aigrader_doc_convert_")
        tmp_dir = Path(self._temp_dir.name)
        completed = subprocess.run(
            [
                converter,
                "--headless",
                "--nologo",
                "--nofirststartwizard",
                "--convert-to",
                "docx",
                "--outdir",
                str(tmp_dir),
                str(self.source_path),
            ],
            capture_output=True,
            text=True,
            timeout=120,
            check=False,
        )
        generated = list(tmp_dir.glob("*.docx"))
        if completed.returncode != 0 or not generated:
            error_text = (completed.stderr or completed.stdout or "转换程序未生成 docx 文件。").strip()
            raise ValueError(f".doc 转 .docx 失败：{error_text[:300]}")
        self.docx_path = generated[0]
        return self.docx_path

    def __exit__(self, exc_type, exc, tb) -> None:
        if self._temp_dir is not None:
            self._temp_dir.cleanup()


def _converted_doc_to_docx(source_path: Path) -> _ConvertedDocContext:
    return _ConvertedDocContext(source_path)


def _find_office_converter() -> str | None:
    # Reuse the converter discovery and bundled LibreOffice extraction used by
    # the high-fidelity Word preview feature.
    try:
        from app.services.docx_preview_service import _find_office_converter as find_converter

        return find_converter()
    except Exception:
        return None


def _parse_pdf_text_with_pypdf(path: Path) -> str:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception as exc:
        raise RuntimeError("解析 PDF 需要安装 PyMuPDF 或 pypdf。请确认运行环境已包含 PDF 解析依赖。") from exc

    try:
        reader = PdfReader(str(path))
    except Exception as exc:
        raise ValueError(f"无法读取 PDF 文件：{exc}") from exc

    sections: list[str] = []
    for page_index, page in enumerate(reader.pages, start=1):
        try:
            text = (page.extract_text() or "").strip()
        except Exception:
            text = ""
        if text:
            sections.append(f"【PDF第{page_index}页】")
            sections.append(text)
    return "\n\n".join(sections).strip()


def _try_load_pymupdf():
    try:
        import fitz  # type: ignore

        return fitz
    except Exception:
        return None
