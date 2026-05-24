from __future__ import annotations

import json
import os
import sys
import tempfile
from io import BytesIO
from pathlib import Path
from typing import Any


PROJECT_ROOT = Path(__file__).resolve().parents[1]
BACKEND_ROOT = PROJECT_ROOT / "backend"


def _set_test_environment(runtime_dir: Path) -> None:
    os.environ["AIGRADER_DATA_DIR"] = str(runtime_dir / "data")
    os.environ["AIGRADER_DB_PATH"] = str(runtime_dir / "data" / "app.db")
    os.environ["AIGRADER_UPLOADS_DIR"] = str(runtime_dir / "uploads")
    os.environ["AIGRADER_EXPORTS_DIR"] = str(runtime_dir / "exports")
    sys.path.insert(0, str(BACKEND_ROOT))


def _assert_ok(response, label: str) -> dict[str, Any]:
    if response.status_code >= 400:
        raise AssertionError(f"{label} HTTP {response.status_code}: {response.text}")
    data = response.json()
    if data.get("ok") is False:
        raise AssertionError(f"{label} returned ok=false: {data}")
    return data


def _create_docx_bytes() -> bytes:
    from docx import Document

    document = Document()
    document.add_heading("实验一 回归测试报告", level=1)
    document.add_paragraph("姓名：张三")
    document.add_paragraph("学号：20260001")
    document.add_paragraph("班级：22测试班")
    document.add_paragraph("实验过程：完成环境配置、运行结果截图整理，并对结果进行了简要分析。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "结果"
    table.cell(1, 0).text = "连通性测试"
    table.cell(1, 1).text = "通过"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _create_ai_scores_for_review(submission_id: int, task_id: int, rubric_items: list[dict[str, Any]]) -> None:
    from app.db.database import execute
    from app.services import scoring_service

    total = 0.0
    for item in rubric_items:
        max_score = float(item["max_score"])
        ai_score = max_score * 0.8
        total += ai_score
        execute(
            """
            INSERT INTO ai_scores (
                submission_id, task_id, rubric_item_id, item_name, max_score,
                ai_score, deduction_reason, evidence_json, suggestion, confidence,
                need_teacher_review
            )
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                submission_id,
                task_id,
                item["id"],
                item["item_name"],
                max_score,
                ai_score,
                "回归测试生成的模拟 AI 扣分原因。",
                json.dumps([], ensure_ascii=False),
                "请教师结合原报告复核。",
                "高",
                1 if item.get("requires_review") else 0,
            ),
        )
    execute(
        """
        UPDATE submissions
        SET ai_status = ?,
            ai_total_score = ?,
            ai_raw_response = ?,
            ai_scored_at = CURRENT_TIMESTAMP,
            updated_at = CURRENT_TIMESTAMP
        WHERE id = ?
        """,
        (scoring_service.AI_DONE, round(total, 2), json.dumps({"source": "smoke_regression"}), submission_id),
    )


def run() -> None:
    with tempfile.TemporaryDirectory(prefix="aigrader_v2_smoke_", ignore_cleanup_errors=True) as tmp_dir_name:
        runtime_dir = Path(tmp_dir_name)
        _set_test_environment(runtime_dir)

        from fastapi.testclient import TestClient
        from app.main import create_app

        with TestClient(create_app()) as client:
            health = _assert_ok(client.get("/api/health"), "health")
            assert health["database"].endswith("app.db")

            course = _assert_ok(
                client.post(
                "/api/courses",
                json={
                    "course_name": "回归测试课程",
                    "course_type": "实验类课程",
                    "semester": "2026-2027第一学期",
                    "description": "自动回归测试临时课程。",
                },
            ),
                "create course",
            )
            course_id = int(course["data"]["id"])

            teaching_class = _assert_ok(
                client.post(
                "/api/classes",
                json={"class_name": "22测试班", "description": "自动回归测试临时班级。", "course_id": course_id},
            ),
                "create class",
            )
            class_id = int(teaching_class["data"]["id"])

            _assert_ok(
                client.post(
                f"/api/students/class-roster?class_id={class_id}",
                json={"student_no": "20260001", "student_name": "张三", "class_name": "22测试班"},
            ),
                "create class student",
            )
            _assert_ok(
                client.post(f"/api/students/course-roster/sync-from-class?course_id={course_id}&class_id={class_id}"),
                "sync course roster",
            )
            course_roster = _assert_ok(
                client.get(f"/api/students/course-roster?course_id={course_id}&class_id={class_id}"),
                "list course roster",
            )
            assert len(course_roster["data"]) == 1

            experiment = _assert_ok(
                client.post(
                "/api/experiments",
                json={
                    "course_id": course_id,
                    "experiment_name": "实验一 自动化回归测试",
                    "experiment_objectives": "验证报告上传、解析、评分复核与导出主链路。",
                    "experiment_requirements": "学生需提交包含姓名、学号、班级、过程说明和结果表格的 docx 报告。",
                    "required_screenshots": "本测试不强制截图。",
                    "key_evaluation_points": "检查报告结构、过程说明和结果分析。",
                    "common_errors": "缺少过程说明；只给结论不做分析。",
                    "special_notes": "自动测试数据，运行结束后会删除。",
                },
            ),
                "create experiment",
            )
            experiment_id = int(experiment["data"]["id"])

            rubric_payloads = [
            {
                "experiment_id": experiment_id,
                "item_name": "报告结构完整性",
                "max_score": 40,
                "description": "检查报告基本结构。",
                "deduction_rules": "缺少关键字段酌情扣分。",
                "requires_review": False,
                "sort_order": 1,
            },
            {
                "experiment_id": experiment_id,
                "item_name": "实验过程与分析",
                "max_score": 60,
                "description": "检查过程记录和结果分析。",
                "deduction_rules": "过程不完整或分析不足酌情扣分。",
                "requires_review": True,
                "sort_order": 2,
            },
        ]
            for payload in rubric_payloads:
                _assert_ok(client.post("/api/rubrics", json=payload), "create rubric")
            rubric_response = _assert_ok(client.get(f"/api/rubrics?experiment_id={experiment_id}"), "list rubrics")
            rubric_items = rubric_response["data"]["items"]
            assert len(rubric_items) == 2
            assert float(rubric_response["data"]["total_score"]) == 100.0

            task = _assert_ok(
                client.post(
                "/api/grading-tasks",
                json={
                    "task_name": "回归测试课程-22测试班-实验一",
                    "course_id": course_id,
                    "class_id": class_id,
                    "experiment_id": experiment_id,
                    "description": "自动回归测试临时批改任务。",
                },
            ),
                "create grading task",
            )
            task_id = int(task["data"]["id"])

            docx_bytes = _create_docx_bytes()
            upload = _assert_ok(
                client.post(
                f"/api/submissions/upload?task_id={task_id}",
                files=[
                    (
                        "files",
                        (
                            "20260001_张三_实验一.docx",
                            docx_bytes,
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        ),
                    )
                ],
            ),
                "upload report",
            )
            assert upload["data"][0]["ok"] is True
            submission_id = int(upload["data"][0]["submission_id"])

            parse_result = _assert_ok(client.post(f"/api/submissions/{submission_id}/parse"), "parse submission")
            assert parse_result["data"]["ok"] is True
            assert int(parse_result["data"]["text_length"]) > 0

            summary = _assert_ok(client.get(f"/api/grading-tasks/{task_id}/summary"), "task summary")
            assert summary["data"]["student_count"] == 1
            assert summary["data"]["upload_count"] == 1

            _create_ai_scores_for_review(submission_id, task_id, rubric_items)
            init_review = _assert_ok(client.post(f"/api/reviews/submissions/{submission_id}/initialize"), "initialize review")
            assert init_review["ok"] is True
            detail = _assert_ok(client.get(f"/api/reviews/submissions/{submission_id}"), "review detail")
            teacher_items = detail["data"]["teacher_items"]
            assert len(teacher_items) == 2

            save_review = _assert_ok(
                client.post(
                f"/api/reviews/submissions/{submission_id}/save",
                json={
                    "score_rows": teacher_items,
                    "overall_comment": "自动回归测试复核通过。",
                    "reviewer_name": "自动测试",
                    "mark_completed": True,
                },
            ),
                "save review",
            )
            assert save_review["ok"] is True

            preview = _assert_ok(client.get(f"/api/exports/preview/{task_id}?mode=reviewed_only"), "export preview")
            assert preview["data"]["grades"]
            analysis = _assert_ok(client.get(f"/api/exports/analysis/{task_id}?mode=reviewed_only"), "analysis preview")
            assert analysis["data"]["summary"]
            assert analysis["data"]["distribution"]

            protected_course_delete = client.delete(f"/api/courses/{course_id}")
            assert protected_course_delete.status_code == 400
            assert "该课程已有后续数据" in protected_course_delete.text

            protected_delete = client.post("/api/submissions/batch-delete", json=[submission_id])
            assert protected_delete.status_code == 400
            assert "已复核报告不能直接删除" in protected_delete.text
            after_delete = _assert_ok(client.get(f"/api/grading-tasks/{task_id}/summary"), "summary after protected delete")
            assert after_delete["data"]["upload_count"] == 1

    print("V2 smoke regression passed: course/class/student/experiment/rubric/task/upload/parse/review/export/delete protection.")


if __name__ == "__main__":
    run()
